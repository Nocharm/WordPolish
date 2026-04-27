"""Outline + StyleSpec → .docx 바이트 (Phase 3: 표 reembed + 이미지 placeholder)."""

import io
import uuid

from docx import Document
from docx.shared import Mm

from app.domain.outline import Block, Outline
from app.domain.style_spec import StyleSpec
from app.renderer.apply_style import apply_paragraph_style
from app.renderer.inject_numbering import renumber
from app.renderer.reembed_raw import reembed_paragraph, reembed_table


def _setup_page(doc, spec: StyleSpec) -> None:
    section = doc.sections[0]
    section.top_margin = Mm(spec.page.margin_top_mm)
    section.bottom_margin = Mm(spec.page.margin_bottom_mm)
    section.left_margin = Mm(spec.page.margin_left_mm)
    section.right_margin = Mm(spec.page.margin_right_mm)


def _add_paragraph_block(doc, block: Block, spec: StyleSpec) -> None:
    para = doc.add_paragraph(block.text or "")
    apply_paragraph_style(para, block.level, spec, alignment_override=block.alignment)


def _add_image_placeholder(doc, block: Block, spec: StyleSpec) -> None:
    text = "[이미지]"
    if block.caption:
        text = f"{text} {block.caption}"
    para = doc.add_paragraph(text)
    apply_paragraph_style(para, 0, spec)


def _add_field_placeholder(doc, block: Block, spec: StyleSpec) -> None:
    text = "[참조 — Phase 4 예정]"
    if block.preview_text:
        text = f"{text} {block.preview_text}"
    para = doc.add_paragraph(text)
    apply_paragraph_style(para, 0, spec)


def _add_caption_paragraph(doc, caption: str, spec: StyleSpec) -> None:
    para = doc.add_paragraph(caption)
    apply_paragraph_style(para, 0, spec)


def render_docx(
    outline: Outline,
    spec: StyleSpec,
    *,
    user_id: uuid.UUID | None = None,
    job_id: uuid.UUID | None = None,
) -> bytes:
    doc = Document()
    _setup_page(doc, spec)
    blocks = renumber(outline.blocks, spec)

    for b in blocks:
        if b.kind == "paragraph":
            if b.raw_xml_ref and user_id is not None and job_id is not None:
                reembed_paragraph(doc, raw_ref=b.raw_xml_ref, user_id=user_id, job_id=job_id)
            else:
                _add_paragraph_block(doc, b, spec)
            continue

        if b.caption:
            _add_caption_paragraph(doc, b.caption, spec)

        if b.kind == "table":
            if b.raw_ref and user_id is not None and job_id is not None:
                reembed_table(doc, raw_ref=b.raw_ref, user_id=user_id, job_id=job_id, spec=spec)
            else:
                para = doc.add_paragraph(b.markdown or "[표 원본 미보존]")
                apply_paragraph_style(para, 0, spec)
            continue

        if b.kind == "image":
            _add_image_placeholder(doc, b, spec)
            continue

        if b.kind == "field":
            _add_field_placeholder(doc, b, spec)
            continue

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
