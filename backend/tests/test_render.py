"""Outline + StyleSpec → .docx round-trip."""

import io
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domain.outline import Block, Outline
from app.domain.style_spec import StyleSpec
from app.renderer.render_docx import render_docx

SEED = Path(__file__).resolve().parent.parent / "app" / "templates_seed" / "report.json"


def _load_default_spec() -> StyleSpec:
    raw = json.loads(SEED.read_text(encoding="utf-8"))
    return StyleSpec.model_validate(raw["spec"])


def _build_outline() -> Outline:
    return Outline(
        job_id="j-1",
        source_filename="x.docx",
        blocks=[
            Block(id="b-1", kind="paragraph", level=1, text="개요", detected_by="word_style"),
            Block(id="b-2", kind="paragraph", level=0, text="본문 sample.", detected_by="word_style"),
            Block(id="b-3", kind="table", level=0, raw_ref="table-0"),
        ],
    )


def test_render_emits_valid_docx():
    spec = _load_default_spec()
    data = render_docx(_build_outline(), spec)
    doc = Document(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs]
    assert any("개요" in t for t in paras)
    assert "본문 sample." in paras
    # 표 placeholder
    assert any("[표는" in t for t in paras)


def test_render_applies_eastasia_font_for_korean():
    spec = _load_default_spec()
    data = render_docx(_build_outline(), spec)
    doc = Document(io.BytesIO(data))
    # 첫 paragraph "1. 개요" (renumber에 의해 "1. " 붙음)
    p = doc.paragraphs[0]
    run = p.runs[0]
    rPr = run._element.find(qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
    assert rFonts is not None
    assert rFonts.get(qn("w:eastAsia")) == "맑은 고딕"
    assert rFonts.get(qn("w:ascii")) == "Arial"
